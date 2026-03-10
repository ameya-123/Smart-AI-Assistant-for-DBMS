from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash
from datetime import datetime
import os
import re
from io import BytesIO
from fpdf import FPDF 
from docx import Document # pip install python-docx

# Import your existing RAG service
from services.rag_service import get_agentic_response 

app = Flask(__name__)
CORS(app, resources={r"/api/*": {"origins": ["http://localhost:5173", "http://localhost:3000"]}})

# --- MySQL Database Configuration ---
app.config['SQLALCHEMY_DATABASE_URI'] = 'mysql+pymysql://root:root@localhost/smart_tutor'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)

# ==========================================
#           DATABASE MODELS
# ==========================================

class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    email = db.Column(db.String(120), unique=True, nullable=False)
    password_hash = db.Column(db.String(256), nullable=False)
    skill_level = db.Column(db.String(50), nullable=False)
    sessions = db.relationship('ChatSession', backref='user', lazy=True)

class ChatSession(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    title = db.Column(db.String(100), nullable=False, default="New Chat")
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    messages = db.relationship('ChatMessage', backref='session', lazy=True, cascade="all, delete-orphan")

    def to_dict(self):
        return {
            'id': self.id,
            'title': self.title,
            'created_at': self.created_at.isoformat()
        }

class ChatMessage(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    session_id = db.Column(db.Integer, db.ForeignKey('chat_session.id'), nullable=False)
    role = db.Column(db.String(10), nullable=False)
    content = db.Column(db.Text, nullable=False)
    timestamp = db.Column(db.DateTime, default=datetime.utcnow)

    def to_dict(self):
        return {
            'id': self.id,
            'role': self.role,
            'content': self.content,
            'timestamp': self.timestamp.isoformat()
        }

with app.app_context():
    db.create_all()

# ==========================================
#                API ROUTES
# ==========================================

@app.route('/api/signup', methods=['POST'])
def signup():
    data = request.json
    if User.query.filter_by(email=data.get('email')).first():
        return jsonify({'error': 'Email already registered'}), 400
    hashed_password = generate_password_hash(data.get('password'))
    new_user = User(name=data.get('name'), email=data.get('email'), password_hash=hashed_password, skill_level=data.get('skillLevel', 'beginner'))
    db.session.add(new_user)
    db.session.commit()
    return jsonify({'message': 'User created successfully!', 'id': new_user.id, 'name': new_user.name, 'skillLevel': new_user.skill_level}), 201

@app.route('/api/signin', methods=['POST'])
def signin():
    data = request.json
    user = User.query.filter_by(email=data.get('email')).first()
    if user and check_password_hash(user.password_hash, data.get('password')):
        return jsonify({'message': 'Login successful', 'id': user.id, 'name': user.name, 'skillLevel': user.skill_level}), 200
    return jsonify({'error': 'Invalid email or password'}), 401

# --- PDF Generation (FIXED WRAPPING & CODE HIGHLIGHTING) ---
@app.route('/api/generate-report', methods=['POST'])
def generate_report():
    try:
        data = request.json
        user_name = data.get('userName', 'Student')
        topic = data.get('topic', 'DBMS Concept')
        raw_content = data.get('content', '')

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        
        # --- HEADER ---
        pdf.set_fill_color(79, 70, 229) 
        pdf.rect(0, 0, 210, 40, 'F')
        pdf.set_text_color(255, 255, 255)
        pdf.set_font("Arial", 'B', 20)
        pdf.cell(0, 15, txt="SMART AI TUTOR", ln=True, align='C')
        pdf.set_font("Arial", 'I', 10)
        pdf.cell(0, 10, txt=f"Academic Lab Report: {topic}", ln=True, align='C')
        
        pdf.set_text_color(0, 0, 0)
        pdf.ln(20)
        pdf.set_font("Arial", 'B', 11)
        pdf.cell(0, 7, txt=f"STUDENT: {user_name.upper()}", ln=True)
        pdf.cell(0, 7, txt=f"DATE: {datetime.now().strftime('%Y-%m-%d')}", ln=True)
        pdf.ln(5)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(10)

        lines = raw_content.split('\n')
        in_code_block = False

        for line in lines:
            # Detect Code Block Toggle
            if line.strip().startswith('```'):
                in_code_block = not in_code_block
                pdf.ln(2)
                continue

            # Clean Markdown Bold/Italic
            clean_line = re.sub(r'\*\*|__|\*|_|`', '', line).strip()
            if not clean_line and not in_code_block: continue

            # A. Render Code Blocks (SQL Highlight)
            if in_code_block:
                pdf.set_fill_color(245, 245, 247) 
                pdf.set_text_color(31, 41, 55)    
                pdf.set_font("Courier", size=10)   
                
                # Draw light gray background for code
                pdf.multi_cell(0, 7, txt=line, border=0, fill=True)
                continue

            # B. Format Section Headers (Aim, Theory, etc.)
            pdf.set_text_color(0, 0, 0)
            if any(key in clean_line.upper() for key in ["AIM:", "THEORY:", "PROCEDURE:", "CONCLUSION:", "NAME:"]):
                pdf.set_font("Arial", 'B', 12)
                pdf.set_text_color(79, 70, 229)
                pdf.ln(4)
                pdf.multi_cell(0, 10, txt=clean_line)
                pdf.set_text_color(0, 0, 0)
                pdf.set_font("Arial", size=11)

            # C. Enhanced Table Rendering
            elif clean_line.startswith('|'):
                cells = [c.strip() for c in clean_line.split('|') if c.strip()]
                if not cells or '---' in clean_line: continue 
                
                col_width = 190 / len(cells)
                pdf.set_font("Arial", 'B', 10) if any(h in clean_line.upper() for h in ["ID", "NAME", "DATE"]) else pdf.set_font("Arial", size=10)
                
                for cell in cells:
                    pdf.cell(col_width, 10, txt=cell, border=1, align='C')
                pdf.ln(10)

            # D. Standard Paragraphs (using multi_cell for wrapping)
            else:
                pdf.set_font("Arial", size=11)
                pdf.multi_cell(0, 7, txt=clean_line.encode('latin-1', 'replace').decode('latin-1'))
                pdf.ln(2)

        if not os.path.exists('reports'): os.makedirs('reports')
        report_filename = f"reports/Report_{datetime.now().strftime('%H%M%S')}.pdf"
        pdf.output(report_filename)
        return send_file(report_filename, as_attachment=True)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# --- Word Generation ---
@app.route('/api/generate-report-word', methods=['POST'])
def generate_report_word():
    try:
        data = request.json
        user_name = data.get('userName', 'Student')
        topic = data.get('topic', 'DBMS Report')
        content = data.get('content', '')

        doc = Document()
        doc.add_heading('SMART AI TUTOR - LAB REPORT', 0)
        doc.add_paragraph(f"Student: {user_name}")
        doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
        doc.add_heading(topic, level=1)

        clean_content = re.sub(r'\*\*|__|\*|_|`', '', content)
        doc.add_paragraph(clean_content)

        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        return send_file(
            file_stream,
            as_attachment=True,
            download_name=f"Report_{user_name}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# --- History & Chat Routes ---
@app.route('/api/sessions/<int:user_id>', methods=['GET'])
def get_user_sessions(user_id):
    sessions = ChatSession.query.filter_by(user_id=user_id).order_by(ChatSession.created_at.desc()).all()
    return jsonify([s.to_dict() for s in sessions]), 200

@app.route('/api/sessions/<int:session_id>/messages', methods=['GET'])
def get_session_messages(session_id):
    session = ChatSession.query.get(session_id)
    if not session: return jsonify({'error': 'Session not found'}), 404
    return jsonify([m.to_dict() for m in session.messages]), 200

@app.route('/api/chat', methods=['POST'])
def chat():
    data = request.json
    user_query = data.get('message')
    user_id = data.get('user_id') 
    skill_level = data.get('skill_level', 'beginner')
    session_id = data.get('session_id') 

    if not user_query or not user_id:
        return jsonify({"error": "Missing message or user_id"}), 400

    try:
        current_session = None
        is_new_session = False
        if session_id:
            current_session = ChatSession.query.get(session_id)
        if not current_session:
            is_new_session = True
            current_session = ChatSession(user_id=user_id, title="New Chat")
            db.session.add(current_session)
            db.session.commit()
            session_id = current_session.id

        user_db_message = ChatMessage(session_id=session_id, role='user', content=user_query)
        db.session.add(user_db_message)
        db.session.commit()

        recent_messages_desc = ChatMessage.query.filter_by(session_id=session_id)\
                                           .filter(ChatMessage.id != user_db_message.id)\
                                           .order_by(ChatMessage.timestamp.desc())\
                                           .limit(10)\
                                           .all()
        recent_messages = recent_messages_desc[::-1]
        
        history_str = ""
        for msg in recent_messages:
            role_prefix = "Human" if msg.role == 'user' else "AI"
            history_str += f"{role_prefix}: {msg.content}\n"

        ai_response_content = get_agentic_response(user_query, history_str, skill_level)
        ai_db_message = ChatMessage(session_id=session_id, role='ai', content=ai_response_content)
        db.session.add(ai_db_message)
        
        if is_new_session:
             current_session.title = (user_query[:30] + '...') if len(user_query) > 30 else user_query

        db.session.commit()

        return jsonify({
            "answer": ai_response_content, 
            "status": "success",
            "session_id": session_id,
            "message_id": ai_db_message.id
        })

    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e), "status": "failed"}), 500

@app.route('/api/sessions/<int:session_id>', methods=['DELETE'])
def delete_session(session_id):
    try:
        session_to_delete = ChatSession.query.get(session_id)
        if not session_to_delete: return jsonify({'error': 'Session not found'}), 404
        db.session.delete(session_to_delete)
        db.session.commit()
        return jsonify({'message': 'Session deleted successfully', 'id': session_id}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

if __name__ == "__main__":
    app.run(debug=True, port=5000)